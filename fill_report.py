"""Fill 专业综合实验报告模板.docx with StudentsManagerSystem project content."""
import docx
from docx.shared import Pt
from copy import deepcopy

TEMPLATE = r'd:\items\project\CS_project\StudentsManagerSystem\docs\专业综合实验报告模板.docx'
OUTPUT = r'd:\items\project\CS_project\StudentsManagerSystem\docs\实验报告-学生管理系统.docx'

doc = docx.Document(TEMPLATE)
paras = doc.paragraphs

def replace_text(para, text):
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for r in para.runs[1:]:
        r.text = ''

def add_para_after(ref, text):
    """Add a new paragraph after ref paragraph and return it."""
    new_p = docx.oxml.OxmlElement('w:p')
    run_elem = docx.oxml.OxmlElement('w:r')
    rPr = docx.oxml.OxmlElement('w:rPr')
    rFonts = docx.oxml.OxmlElement('w:rFonts')
    rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', '宋体')
    rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')
    rPr.append(rFonts)
    sz = docx.oxml.OxmlElement('w:sz')
    sz.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '21')
    rPr.append(sz)
    run_elem.append(rPr)
    t = docx.oxml.OxmlElement('w:t')
    t.text = text
    t.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', 'preserve')
    run_elem.append(t)
    new_p.append(run_elem)
    ref._element.addnext(new_p)
    return docx.text.paragraph.Paragraph(new_p, ref._element.getparent())

# 封面
replace_text(paras[9],  '姓    名：陈小明')
replace_text(paras[10], '学    号：20240101001')
replace_text(paras[11], '专    业：计算机科学与技术')
replace_text(paras[12], '完成日期：2026年6月17日')

# 题目
replace_text(paras[23], '基于 WPF 的学生管理系统设计与实现')

# 摘要
replace_text(paras[25],
    '本系统是一个基于 .NET 8.0 和 WPF 技术开发的桌面学生管理应用。系统采用 SQLite + '
    'Entity Framework Core 作为数据持久化方案，使用 Repository 模式封装数据访问，'
    'Service 层统一处理业务校验与日志记录，前端采用分层架构实现模块化管理。系统覆盖'
    '六大核心功能模块：用户登录与权限控制、学生档案管理（含家庭信息、奖励处分、体检记录）、'
    '学籍管理（含注册、变动、奖助学金、毕业信息）、成绩管理（含课程管理、成绩录入与导入导出）、'
    '基础数据管理（院系、专业、班级字典维护）以及信息查询统计。系统实现了完整的 CRUD 操作、'
    '模糊搜索、排序分页、CSV 导入导出和操作日志记录。通过本项目的开发，充分实践了 C# 桌面应用开发、'
    'EF Core ORM 框架使用、分层架构设计、WPF 界面编程等核心技术能力。')
for i in range(26, 32):
    replace_text(paras[i], '')

# 关键词
replace_text(paras[37], '学生管理系统；WPF；.NET 8.0；SQLite；EF Core；分层架构')
replace_text(paras[38], '')

# 引言
replace_text(paras[55],
    '随着高校招生规模的不断扩大，学生信息的种类和数量日益增长，传统的手工管理方式已难以满足'
    '高效、准确、安全的管理需求。开发一套功能完善的学生管理系统，对于提升学校信息化管理水平、'
    '减轻管理人员工作负担具有重要的现实意义。\n\n本系统基于 .NET 8.0 框架，采用 WPF 桌面应用技术，'
    '结合 SQLite 数据库和 Entity Framework Core ORM 框架，旨在构建一个界面友好、功能完备、'
    '易于扩展的学生管理平台。系统不仅覆盖了学生基本信息管理，还扩展了家庭联系、奖励处分、'
    '学籍变动、成绩管理等多个业务维度，力求满足日常学生管理工作中的各类需求。\n\n本报告将围绕'
    '系统需求分析、系统设计、系统实现和系统测试四个方面，全面阐述该学生管理系统的开发过程和'
    '关键技术要点。')

# 系统需求分析
replace_text(paras[58],
    '2.1 功能需求\n\n（1）用户登录与权限管理：支持管理员和普通用户两种角色登录，'
    '凭据经 SHA256 哈希存储，管理员拥有全部权限，普通用户受限访问基础数据管理。\n\n'
    '（2）学生档案管理：支持学生基本信息的增删改查、模糊搜索、排序分页、CSV 导入导出；'
    '家庭信息、奖励记录、处分记录、体检信息四个子模块独立 CRUD。\n\n'
    '（3）学籍管理：涵盖学籍注册、学籍变动、奖助学金和毕业信息四大类数据的增删改查。\n\n'
    '（4）成绩管理：课程维护、成绩录入修改、学年学期筛选、CSV 批量导入导出。\n\n'
    '（5）基础数据管理：院系、专业、班级字典维护，名称修改联动更新关联数据。\n\n'
    '（6）查询统计：条件组合查询、统计摘要展示、CSV 导出。\n\n（7）日志管理：启动、登录、'
    '关键操作、异常等事件自动记录到日志文件。')

replace_text(paras[59],
    '2.2 非功能需求\n\n（1）性能：页面加载和数据操作响应时间不超过 2 秒。\n\n'
    '（2）可用性：界面简洁美观，采用统一蓝灰色调主题，布局清晰、操作直观。\n\n'
    '（3）可靠性：所有数据库操作包含完善的异常处理机制。\n\n'
    '（4）安全性：密码 SHA256 哈希存储，区分管理员和普通用户的数据访问权限。\n\n'
    '（5）可维护性：代码采用分层架构，Model/Repository/Service/View 职责分离。')

replace_text(paras[60],
    '2.3 数据需求\n\n系统管理 16 张数据表：Students（学生基本信息）、FamilyInfos（家庭信息）、'
    'RewardRecords（奖励记录）、PunishmentRecords（处分记录）、HealthRecords（体检记录）、'
    'StudentRegistrations（学籍注册）、StatusChangeRecords（学籍变动）、ScholarshipInfos'
    '（奖助学金）、GraduationInfos（毕业信息）、Scores（成绩）、Courses（课程）、'
    'Departments（院系）、Majors（专业）、Classes（班级）、Users（用户账号）、'
    'LookupOptions（字典值）。其中 Students 表的学号和身份证号保持业务唯一性，'
    '各子表通过 StudentNo 字段与学生表关联。')

# 系统设计
add_para_after(paras[61],
    '3.1 总体架构\n\n系统采用分层架构：UI 层（WPF XAML + Code-behind）→ Service 业务层'
    '（校验、日志、合并规则）→ Repository 数据访问层（EF Core 封装）→ Data 层（SQLite + DbContext）。'
    '各层职责清晰，上层依赖下层，通过方法调用通信。典型调用链：\n\n'
    'XAML Code-behind → Service → Repository → DbContext → SQLite\n\n'
    '3.2 数据库设计\n\n采用 SQLite 嵌入式数据库，共 16 张表。Students 表为核心表，'
    '通过 StudentNo 与各子表关联。首次运行时 DatabaseInitializer.Initialize() 自动创建'
    '数据库和种子数据（用户账号、示例学生、课程、院系专业班级、字典值等）。\n\n'
    '3.3 模块划分\n\n系统划分为六大功能模块：用户登录与权限、学生档案管理、学籍管理、'
    '成绩管理、基础数据管理、查询统计。各模块界面集中在 Views/ 目录下，业务逻辑在 '
    'Services/ 中实现。')

# 系统实现
add_para_after(paras[62],
    '4.1 登录模块\n\nApp.OnStartup → DatabaseInitializer.Initialize → LoginWindow 显示。'
    '用户输入凭据后，btnLogin_Click 调用 UsersRepository.ValidateCredentials 验证用户名、'
    '密码哈希和启用状态。登录成功记录角色到 App.CurrentUserRole 并打开主窗口。\n\n'
    '4.2 学生档案模块\n\nStudentArchiveView → StudentService → StudentRepository。'
    '使用 DataGrid 展示列表，支持关键字模糊搜索、排序分页。StudentEditWindow 弹窗编辑，'
    '保存前调用 Validate 校验格式（学号 4-20 位字母数字、身份证 18 位、手机 11 位、'
    '姓名 2-50 位）和唯一性（学号、身份证不重复）。家庭/奖励/处分/体检子模块通过 '
    'StudentService 的统一方法保存。\n\n'
    '4.3 成绩模块\n\nScoreView + ScoreEditWindow，通过 ScoreService 操作。'
    '课程下拉从 Courses 表加载，自动计算总分和等级（优秀/良好/中等/及格/不及格）。'
    'CSV 导入按"学号+学年+学期+课程编号"四维合并。\n\n'
    '4.4 基础数据模块\n\nBasicDataView + BasicDataService 管理院系/专业/班级字典。'
    '修改名称时联动更新关联数据，字典项通过 LookupOptions 表动态配置。')

# 系统测试
add_para_after(paras[63],
    '5.1 测试环境\n\n操作系统：Windows 11 | .NET 8.0 | SQLite | xUnit 测试框架\n\n'
    '5.2 单元测试\n\n（1）InputValidatorTests：学号/姓名/身份证/手机号/邮箱/学年/学期/金额校验。\n'
    '（2）AcademicYearHelperTests：学年标准化和显示格式化。\n'
    '（3）CsvExportHelperTests：CSV 值格式化和转义处理。\n'
    '（4）ServiceResultTests：Service 层统一结果工厂方法验证。\n\n'
    '5.3 测试结果\n\n'
    '| 测试模块 | 测试方法 | 测试数据 | 预期 | 实际 |\n'
    '|---------|---------|---------|:---:|:---:|\n'
    '| InputValidator | ValidateStudentNo | "20240001" | True | ✅ |\n'
    '| InputValidator | ValidateName | "张三" | True | ✅ |\n'
    '| InputValidator | ValidateIdCard | "110101199001011234" | True | ✅ |\n'
    '| InputValidator | ValidatePhone | "13800138000" | True | ✅ |\n'
    '| CsvExportHelper | Escape | "a,b" | "\\"a,b\\"" | ✅ |\n'
    '| ServiceResult | Success() | - | Succeeded=true | ✅ |\n'
    '| ServiceResult | Failure("错误") | - | Succeeded=false | ✅ |\n\n'
    '全部测试通过，覆盖了核心工具类和业务基础的关键路径。')

# 总结
add_para_after(paras[64],
    '通过本次课程设计，成功开发了基于 .NET 8.0 + WPF 的学生管理系统。主要收获：\n\n'
    '（1）掌握了 C# 桌面应用从需求分析到测试验证的全流程开发。\n'
    '（2）深入理解了 WPF 框架的 XAML 布局、数据绑定、样式主题和事件路由。\n'
    '（3）掌握了 EF Core ORM 框架的 DbContext 配置、LINQ 查询和种子数据初始化。\n'
    '（4）实践了分层架构设计，实现 Model/Repository/Service/View 的关注点分离。\n'
    '（5）掌握了 SQLite 嵌入式数据库的集成使用。\n'
    '（6）通过 xUnit 编写了单元测试，提高了核心代码的可靠性。\n\n'
    '改进方向：MVVM 架构可进一步深化、权限系统可扩展到按钮级别、'
    '可引入图表控件增强统计分析能力。本次项目达到了课程设计预期目标。')

# 参考文献
add_para_after(paras[66],
    '[1] Microsoft .NET 8.0 文档. https://learn.microsoft.com/zh-cn/dotnet/\n'
    '[2] WPF 文档. https://learn.microsoft.com/zh-cn/dotnet/desktop/wpf/\n'
    '[3] Entity Framework Core 文档. https://learn.microsoft.com/zh-cn/ef/core/\n'
    '[4] SQLite 官方文档. https://www.sqlite.org/docs.html\n'
    '[5] xUnit 测试框架. https://xunit.net/\n'
    '[6] Mark J. Price. C# 12 and .NET 8 – Modern Cross-Platform Development. Packt, 2023.')

doc.save(OUTPUT)
print(f"✅ 报告已生成: {OUTPUT}")
